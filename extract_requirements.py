import docx
import os

def extract_entry_requirements(doc_path):
    try:
        doc = docx.Document(doc_path)
        text = ""
        for para in doc.paragraphs:
            if 'entry requirements' in para.text.lower():
                # Get the paragraph and the next few paragraphs
                idx = doc.paragraphs.index(para)
                for i in range(idx, min(idx + 10, len(doc.paragraphs))):
                    text += doc.paragraphs[i].text + "\n"
                break
        return text if text else "Entry requirements not found in a standard format."
    except Exception as e:
        return f"Error processing {os.path.basename(doc_path)}: {str(e)}"

# Directory containing the Word documents
programs_dir = r"C:\Program Files\Ampps\www\Advanced-learning\Programes"

# Process each document
for filename in os.listdir(programs_dir):
    if filename.endswith('.docx'):
        file_path = os.path.join(programs_dir, filename)
        print(f"\n--- {filename} ---")
        print(extract_entry_requirements(file_path))
        print("-" * 80)
